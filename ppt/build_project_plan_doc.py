from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("AI智能考试平台5-6天冲刺计划书.docx")
CONTACT_SHEET = Path("ui_reference_contact_sheet.jpg")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 35, 45)
GRAY = RGBColor(90, 96, 105)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def style_table(table, header=True):
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if header and r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, bold=True, color=INK, size=9.5)
            else:
                set_cell_shading(cell, WHITE)
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, color=INK, size=9)


def add_para(doc, text="", size=10.5, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 7 if level == 2 else 5)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    set_table_width(table, widths)
    style_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def configure_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    header = sec.header.paragraphs[0]
    header.text = "AI 智能考试与学习分析平台 | 5-6 天冲刺计划"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=GRAY)
    footer = sec.footer.paragraphs[0]
    footer.text = "Qt 前端 + Spring Boot 后端 + MySQL 数据库"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=GRAY)


def cover(doc):
    add_para(doc, "项目冲刺执行方案", size=11, color=BLUE, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "基于 Qt + Spring Boot + MySQL 的 AI 智能考试与学习分析平台", size=22, color=INK, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "5-6 天完成项目、同步完成项目计划说明书与实训答辩 PPT", size=13, color=GRAY, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "建议执行日期：2026 年 7 月 10 日 - 2026 年 7 月 15/16 日（可按实际开始日期顺延）", size=10.5, color=INK, bold=True, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "项目目标不是只做一个能运行的系统，而是做出可演示、可答辩、可讲亮点、能体现个人成长的完整作品。", size=11, color=INK, after=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_table(
        doc,
        ["定位", "最终要交付的东西"],
        [
            ["软件项目", "Qt 桌面端、Spring Boot 服务端、MySQL 数据库、演示数据、可运行演示流程"],
            ["文档材料", "项目计划说明书、需求与设计说明、测试记录、AI 辅助开发记录"],
            ["答辩材料", "10-12 页答辩 PPT、3-5 分钟演示脚本、备用截图和录屏"],
            ["获奖方向", "完整闭环、界面美观、AI 亮点、稳定兜底、表达清晰"],
        ],
        [1.55, 4.95],
    )
    doc.add_page_break()


def project_overview(doc):
    add_heading(doc, "1. 项目最终目标", 1)
    add_para(doc, "项目名称建议：基于大模型增强的 AI 智能考试与学习分析平台。")
    add_para(doc, "核心闭环：教师维护题库并组卷，学生在线考试，系统自动判分，自动形成错题与薄弱知识点分析，再生成个性化学习建议和可导出的学习报告。")
    add_bullets(doc, [
        "前端：Qt Widgets 或 Qt Quick，重点做登录页、角色首页、题库管理、考试页、分析页、报告页。",
        "后端：Spring Boot 提供 REST API，完成用户、题库、试卷、考试记录、错题、AI 分析、日志等服务。",
        "数据库：MySQL 存储用户、题目、试卷、考试、答案、成绩、错题、日志、AI 结果。",
        "AI：优先实现 AI 错题诊断和学习建议；AI 组卷建议作为加分项；所有 AI 功能必须有本地规则兜底。",
        "展示重点：界面要有教育数据驾驶舱效果，答辩时展示完整流程和个人能力提升过程。",
    ])

    add_heading(doc, "2. 先准备什么", 1)
    add_table(
        doc,
        ["类别", "必须准备", "建议标准"],
        [
            ["开发环境", "JDK 17、IntelliJ IDEA、Qt Creator、MySQL 8、Navicat 或 DataGrip、Postman/Apifox", "第一天上午全部装好并截图，避免后期卡环境"],
            ["项目仓库", "backend、qt-client、docs、ppt、demo-data 五个目录", "每天至少提交一次版本，提交信息写清完成模块"],
            ["演示数据", "3 个角色账号、30-60 道题、2 套试卷、5 名学生成绩、若干错题", "答辩前不临时造数据，直接一键导入"],
            ["AI 接口", "准备 API Key、请求地址、模型名、网络失败提示", "AI 不通时自动显示本地分析，不影响演示"],
            ["UI 参考", "当前文件夹 5 张教育后台/学习平台参考图", "采用左侧导航、顶部搜索/用户信息、卡片、图表、浅色主界面"],
            ["文档素材", "需求说明、数据库表、接口表、测试截图、Vibe Coding 记录", "每天晚上补，不要最后一天才写"],
        ],
        [1.1, 3.0, 2.4],
    )

    add_heading(doc, "3. 成功优先级", 1)
    add_table(
        doc,
        ["优先级", "必须完成的内容", "判断标准"],
        [
            ["P0 闭环", "登录、题库、组卷、考试、判分、成绩、错题、学习建议", "不用 AI 也能完整跑通一遍"],
            ["P1 展示", "首页看板、成绩图表、知识点正确率、报告导出", "答辩老师一眼看出项目完整度"],
            ["P2 加分", "AI 错题诊断、AI 学习建议、AI 组卷建议", "能展示大模型价值，并有失败兜底"],
            ["P3 稳定", "日志、异常提示、多线程导出、备用演示数据", "演示不崩、不卡、可恢复"],
            ["P4 获奖", "统一 UI、PPT 逻辑、演示脚本、个人成长总结", "讲得清楚为什么这个项目比普通考试系统更强"],
        ],
        [1.0, 3.0, 2.5],
    )


def architecture(doc):
    add_heading(doc, "4. 推荐系统架构", 1)
    add_table(
        doc,
        ["层次", "技术", "先做什么", "最后做到什么"],
        [
            ["Qt 前端", "Qt Widgets/QSS/Qt Network/Qt Charts/QThread", "先做登录页和主窗口导航，再做表格页，最后做图表与主题", "界面统一、交互顺畅、能调用后端 API"],
            ["Spring Boot 后端", "Spring Web、MyBatis-Plus/JPA、Validation、JWT 可选", "先建实体、Mapper、Service、Controller，再做 AI 和导出", "接口清晰、错误返回统一、可用 Postman 测"],
            ["MySQL 数据库", "MySQL 8、索引、初始化 SQL", "先建核心表和演示数据，再补日志与 AI 结果表", "断电重启后数据仍在，能支撑图表统计"],
            ["AI 服务", "HTTP 调用大模型 API + 本地模板兜底", "先做后端封装接口，不让 Qt 直接保存 Key", "网络失败时自动走本地分析"],
            ["文档/PPT", "Word、PPT、截图、演示脚本", "从第 1 天就同步积累", "最后一天只排版和演练，不重写内容"],
        ],
        [1.2, 1.55, 2.0, 1.75],
    )
    add_para(doc, "Qt 客户端不直接连 MySQL，推荐通过 Spring Boot 接口访问数据。这样架构更像真实项目，答辩时也更容易体现前后端分层能力。", bold=True)

    add_heading(doc, "5. 数据库核心表", 1)
    add_table(
        doc,
        ["表名", "用途", "关键字段"],
        [
            ["users", "三类用户登录与权限", "id, username, password, role, name, class_name, theme"],
            ["questions", "题库", "id, type, content, option_a/b/c/d, answer, difficulty, knowledge_point, score"],
            ["exams", "试卷/考试", "id, title, duration, creator_id, status, created_time"],
            ["exam_questions", "试卷题目关联", "exam_id, question_id, sort_no, score"],
            ["exam_records", "考试成绩记录", "id, exam_id, student_id, score, correct_count, wrong_count, used_time"],
            ["student_answers", "学生作答明细", "record_id, question_id, user_answer, is_correct, score"],
            ["wrong_questions", "错题分析来源", "student_id, question_id, exam_id, knowledge_point, wrong_time"],
            ["study_suggestions", "学习建议和 AI 诊断", "student_id, exam_id, weak_points, suggestion_text, ai_used"],
            ["operation_logs", "考试日志与系统操作", "user_id, action, detail, created_time"],
            ["vibe_records", "AI 辅助开发记录", "stage, prompt_summary, output_summary, improvement, created_time"],
        ],
        [1.4, 1.8, 3.3],
    )

    add_heading(doc, "6. 后端接口先做这些", 1)
    add_table(
        doc,
        ["模块", "接口", "优先级"],
        [
            ["登录", "POST /api/auth/login, GET /api/users/me", "第 1 天必须"],
            ["题库", "GET/POST/PUT/DELETE /api/questions, POST /api/questions/import", "第 2 天必须"],
            ["组卷", "POST /api/exams/generate, GET /api/exams, GET /api/exams/{id}", "第 2-3 天必须"],
            ["考试", "POST /api/exams/{id}/submit, GET /api/records/student/{id}", "第 3 天必须"],
            ["分析", "GET /api/analysis/student/{id}, GET /api/analysis/class/{examId}", "第 4 天必须"],
            ["AI", "POST /api/ai/wrong-diagnosis, POST /api/ai/study-plan, POST /api/ai/paper-advice", "第 4-5 天加分"],
            ["导出", "GET /api/export/score, GET /api/export/report", "第 4-5 天"],
            ["日志", "GET /api/logs, POST /api/logs", "第 5 天"],
        ],
        [1.2, 4.2, 1.1],
    )


def daily_plan(doc):
    add_heading(doc, "7. 6 天推荐计划：稳妥冲奖版", 1)
    add_para(doc, "默认从 2026 年 7 月 10 日开始执行。如果你实际开始日期不同，把 Day 1 顺延即可。每天晚上必须做一次“能否演示”的检查。")

    days = [
        ("Day 1：2026-07-10", "项目搭建 + 需求定稿 + 数据库/API 骨架",
         [
             ["上午", "先建项目结构", "创建 backend、qt-client、docs、ppt、demo-data；Spring Boot 初始化；Qt 新建主工程；MySQL 新建数据库 exam_ai_platform。", "项目能启动，空窗口和后端健康接口能打开"],
             ["下午", "先做数据库", "建 users/questions/exams/exam_questions/exam_records/student_answers/wrong_questions 表，插入演示账号和 30 道题。", "Navicat 能查到数据，后端能读 users 和 questions"],
             ["晚上", "先做登录闭环", "Spring Boot 写登录接口；Qt 登录页调用接口；按角色跳转到学生/教师/管理员主页。", "三类账号能登录，项目说明书写完“需求+架构+数据库初稿”"],
         ]),
        ("Day 2：2026-07-11", "教师端核心：题库管理 + CSV 导入 + 智能组卷",
         [
             ["上午", "先做题库 CRUD", "后端完成题目增删改查、条件筛选；Qt 用 QTableWidget/QTableView 展示题库，做新增/编辑/删除弹窗。", "教师能维护题目，关闭重开数据不丢"],
             ["下午", "再做 CSV 导入", "规定 CSV 模板；后端校验字段；Qt 选择文件上传；返回成功/失败数量。", "错误行不影响正确题目导入"],
             ["晚上", "最后做组卷", "按题型、难度、知识点、数量抽题；保存 exams 和 exam_questions；Qt 显示生成预览。", "生成试卷无重复题，项目说明书补“模块设计+接口表”"],
         ]),
        ("Day 3：2026-07-12", "学生端核心：在线考试 + 自动判分 + 错题记录",
         [
             ["上午", "先做考试列表和加载试题", "学生端显示可参加考试；进入考试页后加载题目、选项、题号导航。", "学生能进入某套试卷并看到全部题"],
             ["下午", "再做答题与倒计时", "Qt 用 QTimer 实现倒计时；切题保存答案；交卷前确认；时间到自动交卷。", "答题状态不乱，自动交卷可触发"],
             ["晚上", "最后做判分", "后端对比标准答案，写 exam_records/student_answers/wrong_questions；Qt 显示分数、正确数、错题入口。", "完整闭环第一次跑通，录 1 段备用演示视频"],
         ]),
        ("Day 4：2026-07-13", "智能分析：成绩统计 + 图表 + 学习建议 + 报告导出",
         [
             ["上午", "先做成绩统计接口", "后端统计平均分、最高分、最低分、及格率、知识点正确率；学生看个人，教师看班级。", "接口数据能支撑图表"],
             ["下午", "再做图表页面", "Qt Charts 做成绩分布柱状图、知识点正确率柱状图/环形图、学生趋势折线图。", "至少 2 种图表来自数据库"],
             ["晚上", "最后做本地学习建议和导出", "按薄弱知识点生成 1 天/3 天复习计划；导出成绩 CSV 和学习报告 TXT/PDF 可选。", "答辩能展示“考试到学习提升”的闭环"],
         ]),
        ("Day 5：2026-07-14", "AI 亮点 + 日志 + 主题 + 多线程 + UI 美化",
         [
             ["上午", "先做 AI 封装", "后端封装 AI 调用；实现错题诊断、学习建议、组卷建议；失败时返回本地模板结果。", "断网也不影响演示"],
             ["下午", "再补完整度功能", "写 operation_logs；管理员日志页；Qt 导出使用 QThread 或异步提示；主题切换保存到 users.theme。", "功能看起来像完整系统"],
             ["晚上", "最后统一 UI", "按参考图做左侧导航、统计卡片、图表卡片、按钮颜色、圆角、间距；整理图标和 QSS。", "登录页、首页、考试页、分析页视觉统一"],
         ]),
        ("Day 6：2026-07-15/16", "测试优化 + 项目计划说明书 + 答辩 PPT + 演示排练",
         [
             ["上午", "先做全流程测试", "按教师、学生、管理员三条路线测试；记录 Bug；补演示数据；准备失败兜底。", "至少 15 条测试记录和截图"],
             ["下午", "再完成文档和 PPT", "整理项目计划说明书、需求设计、数据库设计、接口说明、测试记录、Vibe Coding 记录；PPT 做 10-12 页。", "文档能交，PPT 能讲"],
             ["晚上", "最后演练答辩", "按 3-5 分钟脚本演示：登录、组卷、考试、判分、分析、AI、导出；准备老师提问答案。", "演示顺畅，出现网络/AI失败也能继续"],
         ]),
    ]
    for title, goal, rows in days:
        add_heading(doc, title, 2)
        add_para(doc, goal, bold=True, color=DARK_BLUE)
        add_table(doc, ["时间", "先做什么", "具体做法", "当天验收"], rows, [0.8, 1.25, 3.25, 1.2])

    add_heading(doc, "8. 5 天压缩方案：时间紧时这样砍", 1)
    add_table(
        doc,
        ["天数", "合并方式", "必须保留", "可以弱化"],
        [
            ["Day 1", "环境、数据库、登录、主窗口一次完成", "三角色登录、数据库初始化", "管理员页先只做日志入口"],
            ["Day 2", "题库、CSV、组卷合并", "题库 CRUD、生成试卷", "CSV 可只支持标准模板"],
            ["Day 3", "考试、判分、错题合并", "考试倒计时、交卷、成绩、错题", "题号导航可简单化"],
            ["Day 4", "统计、图表、AI、导出合并", "2 个图表、本地建议、AI 错题诊断", "AI 组卷建议可放到最后"],
            ["Day 5", "UI、测试、文档、PPT 合并", "统一 UI、说明书、PPT、演示脚本", "深色/护眼主题只保留浅色+深色"],
        ],
        [0.8, 2.0, 2.0, 1.7],
    )
    add_para(doc, "压缩到 5 天时，原则是保闭环、保演示、保文档。不要为了多加一个 AI 按钮牺牲稳定性。", bold=True)


def implementation_steps(doc):
    add_heading(doc, "9. 每个模块的具体实现顺序", 1)
    modules = [
        ("登录与角色权限", [
            "先在 MySQL 插入 admin、teacher、student 三类账号。",
            "后端写登录接口，返回 userId、role、name、theme。",
            "Qt 登录按钮点击后发送 POST 请求，成功后保存当前用户信息。",
            "根据 role 打开不同主界面，所有页面顶部显示当前用户。",
            "每次登录成功/失败写 operation_logs。"
        ]),
        ("题库管理", [
            "先完成 questions 表和后端 CRUD。",
            "Qt 教师端先做题目列表，再做新增和编辑弹窗。",
            "筛选条件先做知识点、难度、题型，再做关键词搜索。",
            "删除前弹确认框，成功后刷新表格。",
            "准备 CSV 模板，导入时返回成功数和失败行原因。"
        ]),
        ("智能组卷", [
            "先做普通规则组卷：题型、难度、知识点、数量。",
            "后端查询候选题，随机抽取并用 question_id 去重。",
            "保存 exams 和 exam_questions。",
            "Qt 显示试卷预览，可查看题目和分值。",
            "AI 组卷建议只辅助教师填写规则，不直接替代生成逻辑。"
        ]),
        ("在线考试与判分", [
            "先做考试列表，再进入考试页加载题目。",
            "每道题保存当前选择，切换题号不丢答案。",
            "QTimer 每秒刷新剩余时间，时间到自动提交。",
            "提交后后端判分，写成绩、答题明细和错题表。",
            "结果页显示分数、正确数、错误数、用时和错题入口。"
        ]),
        ("错题分析与学习建议", [
            "先按 student_id 查询 wrong_questions。",
            "按 knowledge_point 聚合错误次数和正确率。",
            "本地模板生成薄弱点说明和复习计划。",
            "AI 可用时把错题、错误答案、正确答案、知识点组成 Prompt。",
            "AI 不可用时显示本地报告，并提示“已使用本地分析”。"
        ]),
        ("报告导出与多线程", [
            "先导出 CSV 成绩表，保证老师能看到全班成绩。",
            "再导出学生学习报告，包含成绩、错题、薄弱点、建议。",
            "Qt 导出时使用线程或异步状态，显示进度条/完成提示。",
            "导出完成后写日志，文件名包含考试名和日期。",
            "答辩时提前准备一个已导出的报告文件作为备用。"
        ]),
    ]
    for title, items in modules:
        add_heading(doc, title, 2)
        add_numbered(doc, items)


def ui_and_docs(doc):
    add_heading(doc, "10. UI 设计安排：参考图怎么落到 Qt", 1)
    add_table(
        doc,
        ["页面", "参考图特征", "Qt 具体做法"],
        [
            ["登录页", "居中卡片、浅色背景、角色选择", "QWidget + QFrame 登录卡片；用户名、密码、角色下拉框；主按钮使用蓝青色"],
            ["主界面", "左侧深色导航、右侧白色内容区", "QStackedWidget 管理页面；左侧按钮带图标和选中状态"],
            ["教师首页", "统计卡片、待处理任务、近期活动", "卡片展示题目数、试卷数、参考人数、平均分；右侧放日志/待批改"],
            ["学生首页", "课程/任务卡片、成绩趋势", "显示可参加考试、最近成绩、AI 建议摘要"],
            ["考试页", "信息清晰、题号导航、倒计时突出", "左侧题目和选项，右侧题号按钮，顶部倒计时红色强调"],
            ["分析页", "柱状图、折线图、圆环图、AI 报告卡片", "Qt Charts + QTextEdit/卡片展示 AI 诊断和复习计划"],
            ["管理页", "表格、筛选、日志", "QTableView 展示日志和用户，顶部放筛选条件"],
        ],
        [1.15, 2.0, 3.35],
    )
    if CONTACT_SHEET.exists():
        add_para(doc, "UI 参考缩略图：左侧导航、浅色卡片、数据图表、任务面板和教育平台风格是本项目的主要视觉方向。", bold=True)
        doc.add_picture(str(CONTACT_SHEET), width=Inches(6.1))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "11. 项目计划说明书同步写法", 1)
    add_table(
        doc,
        ["日期", "当天补充到说明书的内容", "不要遗漏"],
        [
            ["Day 1", "项目背景、目标、技术路线、需求分析、数据库初稿", "写明 Qt + Spring Boot + MySQL 分层架构"],
            ["Day 2", "题库管理、CSV 导入、智能组卷模块设计", "放接口表、页面截图、数据表说明"],
            ["Day 3", "在线考试、判分、错题记录流程", "画完整业务流程图或用文字步骤说明"],
            ["Day 4", "成绩统计、图表、学习建议、报告导出", "说明数据从数据库来，不是静态图"],
            ["Day 5", "AI 模块、主题切换、日志、多线程、UI 优化", "强调 AI 失败兜底和系统稳定性"],
            ["Day 6", "测试记录、总结、创新点、个人能力提升", "至少放 8-12 张关键截图"],
        ],
        [0.9, 3.6, 2.0],
    )

    add_heading(doc, "12. 答辩 PPT 结构", 1)
    add_table(
        doc,
        ["页码", "标题", "内容"],
        [
            ["1", "项目名称与定位", "AI 智能考试与学习分析平台，说明不是普通考试系统"],
            ["2", "问题背景", "传统考试只给分数，缺少诊断和学习提升闭环"],
            ["3", "系统架构", "Qt 前端、Spring Boot 后端、MySQL、AI 服务"],
            ["4", "功能模块", "三角色、题库、组卷、考试、判分、分析、报告"],
            ["5", "核心流程", "教师组卷到学生学习建议的闭环"],
            ["6", "数据库与接口", "核心表和关键接口，体现工程能力"],
            ["7", "界面展示", "登录页、首页、考试页、分析页截图"],
            ["8", "AI 亮点", "错题诊断、学习建议、组卷建议、失败兜底"],
            ["9", "测试与稳定性", "测试用例、日志、多线程导出、异常处理"],
            ["10", "Vibe Coding 记录", "AI 辅助需求、代码、Bug 修复、文档整理"],
            ["11", "项目成果", "完成度、演示数据、可拓展方向"],
            ["12", "总结与答辩", "个人能力提升、后续优化、结束语"],
        ],
        [0.65, 1.55, 4.3],
    )

    add_heading(doc, "13. 答辩演示脚本", 1)
    add_numbered(doc, [
        "先用一句话定位项目：这是一个把在线考试、自动判分、错题诊断和学习建议串起来的 AI 学习分析平台。",
        "教师登录，展示首页数据卡片和题库管理。",
        "新增或筛选题目，演示 CSV 导入结果。",
        "设置规则生成试卷，展示题目不重复和难度/知识点分布。",
        "学生登录，进入考试，答几道题后交卷。",
        "展示自动判分、错题记录、知识点薄弱分析。",
        "点击 AI 错题诊断或学习建议，说明网络失败会使用本地规则兜底。",
        "教师查看班级成绩图表，导出报告。",
        "管理员查看日志，说明系统完整性。",
        "最后展示 PPT 的创新点：完整闭环、AI 加持、界面美化、Vibe Coding 开发记录。"
    ])


def quality_and_growth(doc):
    add_heading(doc, "14. 每天提高代码能力的安排", 1)
    add_table(
        doc,
        ["每天固定动作", "具体做法", "目的"],
        [
            ["30 分钟复盘", "晚上写今天遇到的 3 个问题、原因、解决方法", "把做项目变成能力沉淀"],
            ["1 次代码整理", "删除重复代码，统一命名，补必要注释", "提升工程质量"],
            ["1 次接口测试", "用 Postman/Apifox 测当天新增接口", "避免 Qt 调接口时才发现后端错"],
            ["1 次演示检查", "从登录开始跑一遍当天已有流程", "保证项目一直可演示"],
            ["1 次截图归档", "把关键页面、错误修复、测试结果放进 docs/screenshots", "最后做 PPT 不慌"],
            ["1 次 AI 使用记录", "记录提示词摘要、AI 输出、你如何修改", "答辩时能讲 Vibe Coding 创新"],
        ],
        [1.3, 3.1, 2.1],
    )

    add_heading(doc, "15. 获奖优化清单", 1)
    add_bullets(doc, [
        "演示必须稳定：提前准备本地演示数据、导出文件、关键页面截图和录屏。",
        "界面必须统一：所有按钮、卡片、导航、表格、图表采用同一套 QSS 色彩和间距。",
        "AI 必须讲价值：不要只说接入了大模型，要说它如何把错题变成学习建议。",
        "数据必须真实流动：图表和报告要来自 MySQL 的考试数据，不能写死。",
        "答辩必须有故事：传统考试只给结果，本项目给诊断、建议和改进路径。",
        "代码必须能解释：准备讲清前后端分层、核心表、判分逻辑、AI 兜底逻辑。",
    ])

    add_heading(doc, "16. 最终验收清单", 1)
    add_table(
        doc,
        ["检查项", "达标标准"],
        [
            ["基础闭环", "教师能维护题库、生成试卷；学生能考试；系统能判分并保存成绩"],
            ["智能分析", "能看到错题、薄弱知识点、学习建议"],
            ["AI 功能", "至少完成 AI 错题诊断和 AI 学习建议，失败时本地兜底"],
            ["图表展示", "至少 2 种图表，数据来自 MySQL"],
            ["导出功能", "成绩表和学习报告至少完成一种稳定导出"],
            ["系统完整", "日志、主题、多线程导出至少完成可演示版本"],
            ["UI 效果", "页面风格统一，重点页面有卡片、图表、导航和清晰层级"],
            ["文档材料", "项目计划说明书、测试记录、PPT、演示脚本齐全"],
            ["答辩准备", "能在 3-5 分钟内完整演示，不依赖现场临时造数据"],
        ],
        [1.5, 5.0],
    )


def main():
    doc = Document()
    configure_doc(doc)
    cover(doc)
    project_overview(doc)
    architecture(doc)
    daily_plan(doc)
    implementation_steps(doc)
    ui_and_docs(doc)
    quality_and_growth(doc)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
